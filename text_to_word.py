"""
Chuyển file .txt (OCR kinh sách) sang Word (.docx).
Format giống sách in: header trang, thơ từng dòng, hội thoại, văn xuôi.
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "text"
OUTPUT_DIR = BASE_DIR / "word"

FONT = "Times New Roman"
SZ_BODY = Pt(13)
SZ_TITLE = Pt(20)
SZ_BOOK = Pt(17)
SZ_SECTION = Pt(15)
SZ_SUB = Pt(12)
SZ_META = Pt(11)
SZ_HEADER = Pt(11)

IND_BODY = 0.75
IND_DIALOGUE = 0.5
IND_VERSE_LABEL = 1.0
IND_VERSE_LINE = 2.2

SKIP_LINE = [
    re.compile(r"^Trang này trống", re.I),
    re.compile(r"không có text nào để (extract|trích xuất)", re.I),
    re.compile(r"^\d+\s+\*?Dịch giả", re.I),
    re.compile(r"^Dịch giả.*\d+\s*$", re.I),
    re.compile(r"^\d{1,3}$"),
]

PAGE_HEADER = re.compile(
    r"^(NAM TUY[ỀỂ]N NGỮ LỤC|BỬU TẠNG LUẬN|CỘI NGUỒN TRUYỀN THỪA|KINH PHÁP BẢO ĐÀN)$"
)

SPEAKERS = (
    r"Tăng|Sư|Hòa [Tt]hượng|Thiền [Ss]ư|Đại chúng|Am chủ|Vị am chủ|Triệu Châu|"
    r"Mã Tổ|Phật|Bồ Tát|Vua|Tọa chủ|Ngài|Ông|Bà|Chư vị|Chúng|"
    r"Học nhơn|Học nhân|Lục đại phu|Lục|Thần Sơn|Vân Nham|Động Sơn|"
    r"Thiên Đồng|Huỳnh Bá|Giảng Tăng|Qui Tông|Chân Tịnh|Thủ tọa|Bá"
)
SPEAK_VERBS = (
    r"nói|đáp|hỏi|thưa|kể|đọc|gọi|than|nhắc lại|lớn tiếng|bạch|tụng|rằng|tiếp"
)

DIALOGUE = re.compile(rf"^({SPEAKERS})\s*({SPEAK_VERBS})", re.I)
DIALOGUE_SPLIT = re.compile(rf"(?=({SPEAKERS})\s*(?:{SPEAK_VERBS}))", re.I)

QA_DIALOGUE = re.compile(
    r"^(\*{0,2}\d+\.\s*HỎI\s*:|➤\s*\*{0,2}ĐÁP\s*:|➢\s*\*{0,2}ĐÁP\s*:)",
    re.I,
)

BOOK_TITLE_KW = ("NGỮ LỤC", "NGŨ LỤC", "LUẬN", "KINH", "GÓP NHẶT", "THỪA", "PHẨM", "ĐÀN", "KHOA HỌC", "PHẬT PHÁP")

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\([^)]{1,300}\))")

SECTION_KW = (
    "LỜI DẦU", "LỜI DỊCH", "LỜI NÓI", "MỤC LỤC", "PHỤ LỤC",
    "TRÍCH TRONG", "VÀI HÀNG", "PHẨM ", "HẾT PHẨM", "HẾT",
    "CỦA NGÀI", "DỊCH NGHĨA",
)

META_KW = (
    "NHÀ XUẤT BẢN", "GIÁO HỘI", "THÀNH HỘI", "Địa chỉ", "ĐT:",
    "Trụ sở", "Văn phòng", "TỔ IN", "PL:", "DL:", "HÀ NỘI",
    "TP.HCM", "thực hiện", "Dịch giả", "Việt Dịch", "Từ Ấn",
)

VERSE_LABEL = re.compile(
    r"^[\*\s]*(?:Bình\s*:|Thiên|Kiêm|Chánh|Trung|Ngũ|.*tụng.*):[\*\s]*$",
    re.I,
)


# ── Helpers ──────────────────────────────────────────────────────

def strip_markers(s: str) -> str:
    return s.strip().strip("*").strip()


def cjk_ratio(s: str) -> float:
    if not s:
        return 0.0
    cjk = sum(1 for c in s if "\u4e00" <= c <= "\u9fff")
    return cjk / len(s)


def normalize_book_title(title: str) -> str:
    """Chuẩn hóa tên sách từ OCR (sửa lỗi nhận dạng)."""
    t = re.sub(r"\s+", " ", strip_markers(title)).upper()
    if re.search(r"DUY\s*L[ƯỤ]C", t) and re.search(r"NG[ŨỮU]\s*L[ƯỤ]C", t):
        return "DUY LỰC NGỮ LỤC"
    if "GÓP NHẶT" in t and "PHẬT TỔ" in t:
        return "GÓP NHẶT LỜI PHẬT TỔ VÀ THÁNH HIỀN"
    if "NAM TUY" in t and "NGỮ LỤC" in t:
        return "NAM TUYỀN NGỮ LỤC"
    if "CỘI NGUỒN" in t and "THỪA" in t:
        return "CỘI NGUỒN TRUYỀN THỪA"
    if "BỬU TẠNG" in t and "LUẬN" in t:
        return "BỬU TẠNG LUẬN"
    if "PHÁP BẢO" in t and "ĐÀN" in t:
        return "KINH PHÁP BẢO ĐÀN"
    return t


def is_book_running_title(title: str) -> bool:
    t = title.upper()
    if "NHÀ XUẤT" in t or len(t) < 6:
        return False
    return any(k in t for k in BOOK_TITLE_KW)


def parse_page_marker(s: str) -> Optional[tuple[int, str]]:
    """Nhận dạng chân trang OCR: '12 DUY LỰC NGỮ LỤC' hoặc 'NAM TUYỀN NGỮ LỤC 7'."""
    s = re.sub(r"\s+", " ", strip_markers(s.strip()))
    if not s or "dịch giả" in s.lower():
        return None
    if s.startswith("➤") or any(k in s for k in ("Địa chỉ", "Tp.", "Quận", "P.", "Q.")):
        return None
    if "▪" in s or "..." in s:
        return None

    m = re.match(r"^(\d{1,4})\s+(.+)$", s)
    if m:
        num = int(m.group(1))
        title = normalize_book_title(m.group(2))
        if is_book_running_title(title):
            return num, title

    m = re.match(r"^(.+?)\s+(\d{1,4})$", s)
    if m:
        title = normalize_book_title(m.group(1))
        num = int(m.group(2))
        if is_book_running_title(title):
            return num, title
    return None


def skip(line: str) -> bool:
    s = line.strip()
    if parse_page_marker(s):
        return False
    return bool(s) and any(p.search(s) for p in SKIP_LINE)


def is_page_header_repeat(s: str, line_no: int, upcoming: list) -> bool:
    if line_no <= 25 or not PAGE_HEADER.match(s.strip()):
        return False
    for nxt in upcoming[:4]:
        n = nxt.strip()
        if n in ("CỦA NGÀI TĂNG TRIỆU", "LỜI DỊCH GIẢ") or n.startswith("*(Trích"):
            return False
    return True


def is_meta(s: str) -> bool:
    return any(k in s for k in META_KW)


def is_section(s: str) -> bool:
    if parse_page_marker(s.strip()):
        return False
    s = strip_markers(s)
    if not s or len(s) > 90:
        return False
    if any(k in s.upper() for k in SECTION_KW):
        return True
    if s.startswith("**HẾT"):
        return True
    if s.isupper() and 5 <= len(s) <= 70:
        ratio = sum(c.isalpha() for c in s) / len(s)
        if ratio > 0.55 and not is_meta(s):
            return True
    return False


def is_separator(s: str) -> bool:
    return bool(re.match(r"^[❊\s]+$", s.strip())) and "❊" in s


def is_verse_label(s: str) -> bool:
    t = strip_markers(s)
    if not t:
        return False
    if VERSE_LABEL.match(s.strip()):
        return True
    if t.endswith(":") and len(t) < 70 and not DIALOGUE.match(t):
        if any(k in t for k in ("tụng", "Tụng", "Chánh", "Trung", "Kiêm", "Bình", "Thiên", "Ngũ Vị")):
            return True
        if re.match(r"^Thiên\s", t):
            return True
        if len(t.split()) <= 8 and re.match(r"^[A-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴĐ]", t):
            return True
    if re.search(r"tụng\s*rằng\s*:$", t, re.I):
        return True
    return False


def is_verse_line(s: str) -> bool:
    t = strip_markers(s)
    if not t or len(t) > 58:
        return False
    if s.strip().startswith("*") and s.strip().endswith("*") and "\n" not in s:
        inner = strip_markers(s)
        if len(inner) < 55 and not inner.endswith(":"):
            return True
    if cjk_ratio(t) > 0.45 and len(t) < 50:
        return True
    if len(t) < 48 and t[-1] in ",，.;" and not t.endswith(":"):
        if sum(1 for c in t if c.isalpha()) > 4:
            return True
    return False


def is_commentary_start(s: str) -> bool:
    t = s.strip()
    return t.startswith("*Bình") or t.startswith("*Giảng") or t.startswith("*Viên")


def is_footnote(s: str) -> bool:
    s = s.strip()
    return s.startswith("---") or re.match(r"^\(\d+\)", s)


def is_quote(s: str) -> bool:
    s = s.strip()
    return (s.startswith('*"') or s.startswith('"')) and len(s) > 20


def ends_sentence(s: str) -> bool:
    s = s.rstrip()
    if not s:
        return True
    return s[-1] in ".!?。:;»\"'…)" or s.endswith("...")


def is_continuation(s: str) -> bool:
    if not s:
        return False
    if (DIALOGUE.match(s) or QA_DIALOGUE.match(s) or parse_page_marker(s)
            or is_section(s) or is_verse_label(s) or is_verse_line(s)):
        return False
    return s[0].islower() or s[0] in "([" 


def fmt_run(run, size=SZ_BODY, bold=False, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic


def add_runs(p, text, size=SZ_BODY, bold=False, italic=False):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            fmt_run(r, size, bold=True, italic=italic)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            fmt_run(r, size, bold=bold, italic=True)
        elif part.startswith("(") and part.endswith(")"):
            r = p.add_run(part)
            fmt_run(r, size, italic=True)
        else:
            r = p.add_run(part)
            fmt_run(r, size, bold=bold, italic=italic)


def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=0.0,
             left_indent=0.0, before=0, after=4, spacing=1.35):
    f = p.paragraph_format
    f.alignment = align
    f.first_line_indent = Cm(first_indent) if first_indent else None
    f.left_indent = Cm(left_indent) if left_indent else None
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    f.line_spacing = spacing


def set_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_new_section(doc: Document):
    """Tạo section mới (trang mới) — mỗi section có header riêng."""
    return doc.add_section(WD_SECTION.NEW_PAGE)


def setup_section_header(section, title: str, page_num: int):
    """Header giống sách in: tên trái, số trang phải (số cố định từ OCR)."""
    section.header.is_linked_to_previous = False
    hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    hp.clear()
    para_fmt(hp, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    r1 = hp.add_run(title.upper())
    fmt_run(r1, SZ_HEADER, bold=True)
    hp.add_run("\t")
    r2 = hp.add_run(str(page_num))
    fmt_run(r2, SZ_HEADER, bold=True)
    set_bottom_border(hp)


# ── Parse ────────────────────────────────────────────────────────

def next_nonempty(raw_lines, start):
    for j in range(start, len(raw_lines)):
        if raw_lines[j].strip():
            return raw_lines[j].strip()
    return ""


def clean_and_join(raw: str) -> list[str]:
    all_lines = raw.splitlines()
    raw_lines = []
    for no, ln in enumerate(all_lines, 1):
        s = ln.strip()
        if not s:
            raw_lines.append("")
            continue
        upcoming = [all_lines[j].strip() for j in range(no, min(no + 4, len(all_lines)))]
        if skip(s) or is_page_header_repeat(s, no, upcoming):
            continue
        raw_lines.append(ln.rstrip())

    merged, buf = [], ""
    verse_mode = False

    for idx, line in enumerate(raw_lines):
        s = line.strip()
        if not s:
            nxt = next_nonempty(raw_lines, idx + 1)
            if buf and nxt and is_continuation(nxt) and not verse_mode:
                continue
            if buf:
                merged.append(buf)
                buf = ""
            verse_mode = False
            continue

        if is_verse_label(s):
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(s)
            verse_mode = True
            continue

        if is_verse_line(s):
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(s)
            verse_mode = True
            continue

        if verse_mode and len(s) > 65:
            verse_mode = False

        marker = parse_page_marker(s)
        if marker:
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(s)
            verse_mode = False
            continue

        if (is_separator(s) or is_section(s) or is_commentary_start(s)
                or is_footnote(s) or DIALOGUE.match(s) or QA_DIALOGUE.match(s)
                or is_quote(s)
                or s.startswith("Địa chỉ") or s.startswith("ĐT:")
                or (is_meta(s) and len(s) < 80)
                or (s.isupper() and len(s) < 55)
                or s.startswith("HÀ NỘI")):
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(s)
            verse_mode = False
            continue

        if "tụng rằng:" in s.lower() or "bài tụng rằng:" in s.lower():
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(s)
            verse_mode = True
            continue

        if buf:
            if not ends_sentence(buf) or is_continuation(s):
                buf += " " + s
            else:
                merged.append(buf)
                buf = s
        else:
            buf = s

    if buf:
        merged.append(buf)
    return merged


def split_dialogues(text: str) -> list[str]:
    if not DIALOGUE_SPLIT.search(text):
        return [text]
    parts, last = [], 0
    for m in DIALOGUE_SPLIT.finditer(text):
        pos = m.start()
        if pos == 0:
            continue
        before = text[:pos].rstrip()
        if not before or before[-1] not in '.!?"»\'':
            continue
        chunk = text[last:pos].strip()
        if chunk:
            parts.append(chunk)
        last = pos
    tail = text[last:].strip()
    if tail:
        parts.append(tail)
    return parts if len(parts) > 1 else [text]


def parse_blocks(lines: list[str]) -> tuple[list[dict], str]:
    blocks = []
    book_title = ""
    i, n = 0, len(lines)
    cover_done = False

    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue

        if not cover_done:
            cover_lines = []
            while i < n and len(cover_lines) < 15:
                s = lines[i].strip()
                if not s:
                    i += 1
                    if cover_lines:
                        break
                    continue
                if s == "NAM TUYỀN NGỮ LỤC" and len(cover_lines) >= 3:
                    break
                if len(s) > 60 and not is_meta(s) and not s.isupper():
                    break
                cover_lines.append(s)
                i += 1
            if cover_lines:
                blocks.append({"type": "cover", "lines": cover_lines})
                for ln in cover_lines:
                    if (ln.isupper() and len(ln) > 10 and not is_meta(ln)
                            and "THÍCH" not in ln and "GIÁO HỘI" not in ln
                            and "THÀNH HỘI" not in ln and "TỔ IN" not in ln):
                        book_title = ln
                        break
                cover_done = True
            continue

        marker = parse_page_marker(s)
        if marker:
            num, title = marker
            blocks.append({"type": "new_page", "page_num": num, "title": title})
            i += 1
            continue

        if is_separator(s):
            blocks.append({"type": "sep"})
            i += 1
            continue

        if is_verse_label(s):
            blocks.append({"type": "verse_label", "text": strip_markers(s)})
            i += 1
            continue

        if is_verse_line(s):
            blocks.append({"type": "verse_line", "text": strip_markers(s)})
            i += 1
            continue

        if is_commentary_start(s):
            inner = strip_markers(s)
            if is_verse_label(s) or (inner.endswith(":") and len(inner) < 80):
                blocks.append({"type": "verse_label", "text": inner})
            elif len(inner) < 60 and not inner.endswith("."):
                blocks.append({"type": "verse_line", "text": inner, "italic": True})
            else:
                blocks.append({"type": "commentary", "text": inner})
            i += 1
            while i < n and lines[i].strip().startswith("*"):
                t = lines[i].strip()
                inner = strip_markers(t)
                if is_verse_line(t) or (len(inner) < 55 and not inner.endswith(".")):
                    blocks.append({"type": "verse_line", "text": inner, "italic": True})
                else:
                    blocks.append({"type": "commentary", "text": inner})
                i += 1
            continue

        if is_footnote(s):
            fn = [s.lstrip("-").strip()]
            i += 1
            while i < n and lines[i].strip() and not lines[i].strip().startswith("---"):
                if is_section(lines[i]) or DIALOGUE.match(lines[i].strip()):
                    break
                fn.append(lines[i].strip())
                i += 1
            blocks.append({"type": "footnote", "text": " ".join(fn)})
            continue

        if is_quote(s):
            blocks.append({"type": "quote", "text": strip_markers(s)})
            i += 1
            continue

        if is_section(s):
            sec_lines = [s]
            i += 1
            while i < n:
                nxt = lines[i].strip()
                if not nxt:
                    break
                if parse_page_marker(nxt):
                    break
                if nxt.isupper() and len(nxt) < 50 and not DIALOGUE.match(nxt) and not is_meta(nxt):
                    sec_lines.append(nxt)
                    i += 1
                elif nxt.startswith("*(") or nxt.startswith("*(Trích"):
                    sec_lines.append(nxt)
                    i += 1
                    break
                else:
                    break
            if sec_lines[0].strip("*") == "**HẾT**" and "BỬU TẠNG LUẬN" not in sec_lines:
                sec_lines.insert(1, "BỬU TẠNG LUẬN")
            blocks.append({"type": "section", "lines": sec_lines})
            if not book_title and sec_lines[0].isupper():
                book_title = strip_markers(sec_lines[0])
            continue

        text = s
        i += 1
        while i < n and lines[i].strip():
            nxt = lines[i].strip()
            if (parse_page_marker(nxt) or is_separator(nxt) or is_section(nxt)
                    or is_commentary_start(nxt) or is_footnote(nxt)
                    or is_verse_label(nxt) or is_verse_line(nxt)
                    or DIALOGUE.match(nxt) or QA_DIALOGUE.match(nxt)):
                break
            if ends_sentence(text):
                break
            text += " " + nxt
            i += 1

        text = re.sub(r"\s+", " ", text).strip()
        if not text:
            continue

        parts = split_dialogues(text)
        if len(parts) > 1 or DIALOGUE.match(text) or QA_DIALOGUE.match(text):
            for part in parts:
                blocks.append({
                    "type": "dialogue" if (DIALOGUE.match(part) or QA_DIALOGUE.match(part)) else "body",
                    "text": part,
                })
        else:
            blocks.append({"type": "body", "text": text})

    return blocks, book_title or "KINH SÁCH"


# ── Render ───────────────────────────────────────────────────────

def setup_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.8)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.0)
    sec.different_first_page_header_footer = True

    st = doc.styles["Normal"]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = SZ_BODY


def render_cover(doc, lines: list[str]):
    for ln in lines:
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        if ln in ("VÀ", "và"):
            add_runs(p, ln, SZ_BOOK)
        elif ln.startswith("Địa chỉ"):
            add_runs(p, ln.split("ĐT:")[0].strip(), SZ_META)
            if "ĐT:" in ln:
                p2 = doc.add_paragraph()
                para_fmt(p2, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
                add_runs(p2, "ĐT:" + ln.split("ĐT:", 1)[1], SZ_META)
            continue
        elif ln.startswith("ĐT:"):
            add_runs(p, ln, SZ_META)
        elif is_meta(ln):
            add_runs(p, ln, SZ_META)
        elif ln.startswith("*") and ln.endswith("*"):
            add_runs(p, ln, SZ_META, italic=True)
        elif ln.isupper() and len(ln) > 10:
            sz = SZ_TITLE if ln in ("NAM TUYỀN NGỮ LỤC", "BỬU TẠNG LUẬN") else SZ_SECTION
            add_runs(p, ln, sz, bold=True)
        elif ln.isupper():
            add_runs(p, ln, SZ_BOOK, bold=True)
        else:
            add_runs(p, ln, SZ_META)
    sec = add_new_section(doc)
    sec.header.is_linked_to_previous = False
    if sec.header.paragraphs:
        sec.header.paragraphs[0].clear()


def render_section(doc, lines: list[str]):
    text = lines[0].strip()
    if text.strip("*") == "BỬU TẠNG LUẬN" and any(
        "CỦA NGÀI" in ln or "LỜI DỊCH" in ln for ln in lines[1:]
    ):
        doc.add_page_break()

    for j, ln in enumerate(lines):
        p = doc.add_paragraph()
        clean = strip_markers(ln)
        if j == 0 and clean.isupper() and len(clean) > 8:
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=8, spacing=1.2)
            add_runs(p, clean, SZ_TITLE if clean == "BỬU TẠNG LUẬN" else SZ_SECTION, bold=True)
        elif ln.startswith("*(") or ln.startswith("**"):
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
            add_runs(p, ln, SZ_SUB, italic=ln.startswith("*("))
        elif clean.startswith("HẾT"):
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=20, after=20)
            add_runs(p, clean, SZ_SECTION, bold=True)
        elif clean.isupper():
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6)
            add_runs(p, clean, SZ_SECTION, bold=True)
        else:
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
            add_runs(p, clean, SZ_BOOK, bold=True)


def render_new_page(doc: Document, page_num: int, title: str, default_title: str):
    sec = add_new_section(doc)
    setup_section_header(sec, title or default_title, page_num)


def render_blocks(doc, blocks: list[dict], default_title: str):
    for blk in blocks:
        t = blk["type"]

        if t == "cover":
            render_cover(doc, blk["lines"])

        elif t == "new_page":
            render_new_page(doc, blk["page_num"], blk["title"], default_title)

        elif t == "section":
            render_section(doc, blk["lines"])

        elif t == "sep":
            p = doc.add_paragraph()
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, spacing=1.0)
            add_runs(p, "❊    ❊    ❊", SZ_META)

        elif t == "verse_label":
            p = doc.add_paragraph()
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=IND_VERSE_LABEL,
                     after=2, spacing=1.25)
            add_runs(p, blk["text"], SZ_BODY, bold=True)

        elif t == "verse_line":
            p = doc.add_paragraph()
            para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=IND_VERSE_LINE,
                     after=2, spacing=1.2)
            add_runs(p, blk["text"], SZ_BODY, italic=blk.get("italic", False))

        elif t == "commentary":
            p = doc.add_paragraph()
            para_fmt(p, left_indent=IND_VERSE_LABEL, after=4, spacing=1.25)
            add_runs(p, blk["text"], SZ_BODY, italic=True)

        elif t == "footnote":
            p = doc.add_paragraph()
            para_fmt(p, left_indent=0.5, after=4, spacing=1.2)
            add_runs(p, blk["text"], SZ_META, italic=True)

        elif t == "quote":
            p = doc.add_paragraph()
            para_fmt(p, left_indent=1.0, after=6, spacing=1.3)
            add_runs(p, blk["text"], SZ_BODY, italic=True)

        elif t == "dialogue":
            p = doc.add_paragraph()
            para_fmt(p, left_indent=IND_DIALOGUE, after=4)
            m = DIALOGUE.match(blk["text"])
            qa = QA_DIALOGUE.match(blk["text"]) if not m else None
            if m:
                speaker = blk["text"][: m.end()].strip()
                rest = blk["text"][m.end() :].strip().lstrip(":").strip()
                r = p.add_run(speaker + ": ")
                fmt_run(r, SZ_BODY, bold=True)
                add_runs(p, rest, SZ_BODY)
            elif qa:
                label = strip_markers(blk["text"][: qa.end()]).strip()
                rest = blk["text"][qa.end() :].strip().lstrip(":").strip()
                r = p.add_run(label + " ")
                fmt_run(r, SZ_BODY, bold=True)
                add_runs(p, rest, SZ_BODY)
            else:
                add_runs(p, blk["text"])

        elif t == "body":
            p = doc.add_paragraph()
            para_fmt(p, first_indent=IND_BODY, after=5)
            add_runs(p, blk["text"])


def text_to_docx(txt_path: Path, docx_path: Path):
    raw = txt_path.read_text(encoding="utf-8")
    lines = clean_and_join(raw)
    blocks, book_title = parse_blocks(lines)

    doc = Document()
    setup_doc(doc)
    render_blocks(doc, blocks, book_title)
    doc.save(docx_path)


def xu_ly_tat_ca(input_dir: Path, output_dir: Path, only: Optional[str] = None):
    output_dir.mkdir(exist_ok=True)
    files = sorted(input_dir.glob("*.txt"))
    if only:
        files = [f for f in files if f.stem == only]

    if not files:
        print("⚠️  Không có file .txt")
        return

    print(f"Chuyển {len(files)} file → Word\n")
    for f in files:
        out = output_dir / f"{f.stem}.docx"
        print(f"📝 {f.name} → {out.name}")
        try:
            text_to_docx(f, out)
            print(f"   ✅ {out.stat().st_size // 1024} KB")
        except Exception as e:
            print(f"   ❌ {e}")
    print(f"\n🎉 Xong! → {output_dir}")


if __name__ == "__main__":
    import sys

    only = sys.argv[1] if len(sys.argv) > 1 else None
    xu_ly_tat_ca(INPUT_DIR, OUTPUT_DIR, only=only)
